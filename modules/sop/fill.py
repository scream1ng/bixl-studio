"""
fill_sop.py — Fill IXL SOP Word template with step images and text.

Entry point:
    fill_sop(steps, part_no, part_name, doc_no,
             format_key="a3-landscape", steps_per_page=8,
             template_path=None, output_path=None)

steps is a list of dicts: {"image": path|bytes, "text": str}

Template selection:
  Picks from docx/sop/templates/{format_key}__{steps_per_page}steps.docx.
  Pass template_path to override.

Template layout:
  Each page table repeats (heading / image / text) blocks.
  Grid (cols x n_blocks) is inferred from the table dimensions.
  Row pattern per block i: heading=i*3, image=i*3+1, text=i*3+2.

Multi-page:
  Steps beyond template capacity: deep-copy template table after page break,
  continue step numbering, blank headings for unfilled trailing slots.
"""
from __future__ import annotations

import copy
import io
import os
from typing import Optional, Union

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage

ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
TEMPLATE_DIR = os.path.join(ROOT, "docx", "sop", "templates")

IMG_PADDING_CM = 0.18   # vertical breathing room inside image cell
BODY_FONT = "Century Gothic"
BODY_PT = 11


# Template helpers

def _template_path(format_key: str, steps_per_page: int) -> str:
    return os.path.join(TEMPLATE_DIR, f"{format_key}__{steps_per_page}steps.docx")


def _infer_grid(table) -> tuple[int, int]:
    """Return (cols, n_blocks) from a fill-template table."""
    return len(table.columns), len(table.rows) // 3


def _col_width_cm(table, col: int) -> float:
    tc = table.rows[0].cells[col]._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is not None:
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is not None:
            return int(tcW.get(qn("w:w"))) / 1440 * 2.54
    return 12.0


def _image_row_height_cm(table, block: int) -> float:
    trPr = table.rows[block * 3 + 1]._tr.find(qn("w:trPr"))
    if trPr is not None:
        trH = trPr.find(qn("w:trHeight"))
        if trH is not None:
            return int(trH.get(qn("w:val"))) / 1440 * 2.54
    return 5.5


# Cell helpers

def _clear_cell(cell) -> None:
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for run in list(cell.paragraphs[0].runs):
        run._element.getparent().remove(run._element)


def _set_cell_vertical_align(cell, align: str = "center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _open_image(image: Union[str, bytes]) -> PILImage.Image:
    if isinstance(image, (bytes, bytearray)):
        return PILImage.open(io.BytesIO(image))
    return PILImage.open(image)


def _image_source(image: Union[str, bytes]):
    if isinstance(image, (bytes, bytearray)):
        return io.BytesIO(image)
    return image


def add_image_to_cell(cell, image: Union[str, bytes],
                      max_width_cm: float = 12.0,
                      max_height_cm: float = 5.5) -> None:
    """Insert image scaled to fit cell, horizontally + vertically centred."""
    img = _open_image(image)
    w, h = img.size
    scale = min(max_width_cm / w, max_height_cm / h)
    final_width_cm = w * scale

    _clear_cell(cell)
    _set_cell_vertical_align(cell, "center")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(_image_source(image), width=Cm(final_width_cm))


def add_text_to_cell(cell, text: str) -> None:
    _clear_cell(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text or "")
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_PT)


def set_heading_text(cell, text: str) -> None:
    """Set or blank STEP N heading run, preserving existing formatting."""
    para = cell.paragraphs[0]
    if para.runs:
        first = para.runs[0]
        first.text = text
        for extra in list(para.runs[1:]):
            extra._element.getparent().remove(extra._element)
    elif text:
        run = para.add_run(text)
        run.font.name = BODY_FONT


# Header / footer field replacement (handles split runs)

def _replace_paragraph_text(para, new_text: str) -> None:
    if not para.runs:
        para.add_run(new_text)
        return
    first = para.runs[0]
    rPr = first._r.find(qn("w:rPr"))
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(copy.deepcopy(rPr))
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = new_text
    new_r.append(new_t)
    para._p.append(new_r)


def set_part_title(doc, part_no: str, part_name: str) -> None:
    """Header table row 1 cell 2: '<part_no> - <part_name>'."""
    header = doc.sections[0].header
    if not header.tables:
        return
    para = header.tables[0].rows[1].cells[2].paragraphs[0]
    title = f"{part_no} - {part_name}" if part_no else part_name
    _replace_paragraph_text(para, title)


def set_doc_number(doc, doc_no: str) -> None:
    """Footer table row 0 cell 0: '<doc_no>'."""
    footer = doc.sections[0].footer
    if not footer.tables:
        return
    para = footer.tables[0].rows[0].cells[0].paragraphs[0]
    _replace_paragraph_text(para, doc_no)


# Page table fill

def _fill_page_table(table, page_steps: list, base_index: int) -> None:
    """Fill one template table. Blank heading for any unfilled trailing slot."""
    cols, n_blocks = _infer_grid(table)
    steps_per_page = cols * n_blocks

    for slot in range(steps_per_page):
        block = slot // cols
        col = slot % cols
        step_no = base_index + slot + 1

        heading_cell = table.rows[block * 3].cells[col]
        image_cell = table.rows[block * 3 + 1].cells[col]
        text_cell = table.rows[block * 3 + 2].cells[col]

        step = page_steps[slot] if slot < len(page_steps) else None
        has_content = bool(step and (step.get("image") or step.get("text")))

        if has_content:
            set_heading_text(heading_cell, f"STEP {step_no}")
            if step.get("image"):
                cw = _col_width_cm(table, col) - 0.2
                rh = _image_row_height_cm(table, block) - IMG_PADDING_CM * 2
                add_image_to_cell(image_cell, step["image"],
                                  max_width_cm=cw, max_height_cm=rh)
            if step.get("text"):
                add_text_to_cell(text_cell, step["text"])
        else:
            # Empty step: clear heading, image, and text so copied content doesn't bleed through.

            set_heading_text(heading_cell, "")
            _clear_cell(image_cell)
            _clear_cell(text_cell)


def _append_page_break_and_table(doc, source_table):
    """Deep-copy source_table into body before sectPr, force it onto a new page."""
    body = doc.element.body
    children = list(body)
    sectPr = body.find(qn("w:sectPr"))
    insert_idx = children.index(sectPr) if sectPr is not None else len(children)

    new_tbl = copy.deepcopy(source_table._tbl)
    body.insert(insert_idx, new_tbl)
    from docx.table import Table
    table = Table(new_tbl, doc)

    # Force page break before this table via the first heading cell's paragraph.
    # A separate break paragraph can strand on the last page and push the table
    # one page further than expected.
    first_para = table.rows[0].cells[0].paragraphs[0]
    pPr = first_para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:pageBreakBefore")):
        pPr.remove(old)
    pPr.append(OxmlElement("w:pageBreakBefore"))
    return table


# Public entry point

def fill_sop(steps,
             part_no: str,
             part_name: str,
             doc_no: str,
             format_key: str = "a3-landscape",
             steps_per_page: int = 8,
             template_path: Optional[str] = None,
             output_path: Optional[str] = None) -> Union[str, bytes]:
    """Fill the SOP template and return bytes or save to output_path.

    Args:
        steps: list of {"image": path|bytes, "text": str}, any length.
        part_no, part_name, doc_no: header/footer values.
        format_key: "a3-landscape", "a3-portrait", "a4-landscape", "a4-portrait".
        steps_per_page: 1, 2, 4, 6, or 8 (must match an available template).
        template_path: override auto-selected template.
        output_path: save location; if None returns bytes.
    """
    if template_path is None:
        template_path = _template_path(format_key, steps_per_page)

    doc = Document(template_path)
    first_table = doc.tables[0]
    cols, n_blocks = _infer_grid(first_table)
    spp = cols * n_blocks

    _fill_page_table(first_table, steps[:spp], base_index=0)

    page = 1
    while page * spp < len(steps):
        start = page * spp
        page_steps = steps[start:start + spp]
        new_table = _append_page_break_and_table(doc, first_table)
        _fill_page_table(new_table, page_steps, base_index=start)
        page += 1

    set_part_title(doc, part_no, part_name)
    set_doc_number(doc, doc_no)

    if output_path:
        doc.save(output_path)
        return output_path

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    import tempfile
    from PIL import Image, ImageDraw

    tmp = tempfile.mkdtemp()
    demo_steps = []
    for i in range(6):
        p = os.path.join(tmp, f"s{i}.jpg")
        im = Image.new("RGB", (400, 300), (60 + i * 25, 90, 120))
        ImageDraw.Draw(im).text((180, 140), f"Step {i + 1}", fill="white")
        im.save(p)
        demo_steps.append({"image": p, "text": f"Description for step {i + 1}."})

    out = os.path.join(tmp, "demo_output.docx")
    fill_sop(demo_steps, "36611", "Tastic Luminate Heat Module", "A0866",
             format_key="a3-landscape", steps_per_page=8, output_path=out)
    print("Wrote", out)
