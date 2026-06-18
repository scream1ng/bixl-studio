"""
fill_sop.py — Fill the IXL SOP A3 Word template with step images and text.

This is the core document engine for BIXL Studio. It opens the real IXL
template (preserving border, logo, PPE icons, header and footer) and injects
an image + description into the correct table cells for each step.

Entry point:
    fill_sop(steps, part_no, part_name, doc_no, template_path, output_path)

`steps` is a list of dicts of any length:
    {"image": <path or bytes>, "text": "description"}

Behaviour:
  - 8 steps per page (two blocks of four).
  - Unused cells on a page are left blank AND their "STEP N" heading is cleared.
  - More than 8 steps -> additional pages are appended (STEP 9, 10, ...),
    each a duplicate of the template table after a page break.
  - Images scale to fit the cell by the tighter of width/height so they never
    overflow the row onto a second page; portrait and landscape both work.
  - Images are vertically centred in their cell.
  - Header part number/name and footer doc number are replaced even though the
    template stores them split across multiple runs.

See CLAUDE.md for the full template structure and rules.
"""

from __future__ import annotations

import copy
import io
import os
from typing import Optional, Union

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image as PILImage


# ── Template layout constants ───────────────────────────────────────────────
# The body table has 6 rows. Steps map to cells as follows (per page):
#   Row 0: STEP 1-4 headings      Row 3: STEP 5-8 headings
#   Row 1: image cells 1-4        Row 4: image cells 5-8
#   Row 2: text cells 1-4         Row 5: text cells 5-8
HEADING_ROWS = (0, 3)
IMAGE_ROWS = (1, 4)
TEXT_ROWS = (2, 5)
STEPS_PER_PAGE = 8
COLS = 4

IMG_MAX_WIDTH_CM = 12.5
IMG_MAX_HEIGHT_CM = 5.5
BODY_FONT = "Century Gothic"
BODY_PT = 11


# ── Cell helpers ────────────────────────────────────────────────────────────

def _clear_cell(cell) -> None:
    """Remove every paragraph except the first, and empty the first's runs."""
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for run in list(cell.paragraphs[0].runs):
        run._element.getparent().remove(run._element)


def _set_cell_vertical_align(cell, align: str = "center") -> None:
    """Set vertical alignment on a cell via its tcPr XML."""
    tcPr = cell._tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(existing)
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def _open_image(image: Union[str, bytes]) -> PILImage.Image:
    if isinstance(image, (bytes, bytearray)):
        return PILImage.open(io.BytesIO(image))
    return PILImage.open(image)


def _image_source(image: Union[str, bytes]):
    """Return something python-docx add_picture accepts (path or BytesIO)."""
    if isinstance(image, (bytes, bytearray)):
        return io.BytesIO(image)
    return image


def add_image_to_cell(cell, image: Union[str, bytes],
                      max_width_cm: float = IMG_MAX_WIDTH_CM,
                      max_height_cm: float = IMG_MAX_HEIGHT_CM) -> None:
    """Insert an image scaled to fit the cell, horizontally + vertically centred."""
    img = _open_image(image)
    w, h = img.size
    scale = min(max_width_cm / w, max_height_cm / h)
    final_width_cm = w * scale

    _clear_cell(cell)
    _set_cell_vertical_align(cell, "center")
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(_image_source(image), width=Cm(final_width_cm))


def add_text_to_cell(cell, text: str) -> None:
    """Insert a description into a cell with the template body font."""
    _clear_cell(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text or "")
    run.font.name = BODY_FONT
    run.font.size = Pt(BODY_PT)


def set_heading_text(cell, text: str) -> None:
    """Set or blank the 'STEP N' heading run, preserving its formatting."""
    para = cell.paragraphs[0]
    if para.runs:
        first = para.runs[0]
        first.text = text
        for extra in list(para.runs[1:]):
            extra._element.getparent().remove(extra._element)
    elif text:
        run = para.add_run(text)
        run.font.name = BODY_FONT


# ── Header / footer field replacement (handles split runs) ──────────────────

def _replace_paragraph_text(para, new_text: str) -> None:
    """Replace all runs in a paragraph with one run, keeping the first run's format."""
    if not para.runs:
        run = para.add_run(new_text)
        return
    first = para.runs[0]
    rPr = first._r.find(qn("w:rPr"))
    for run in list(para.runs):
        run._element.getparent().remove(run._element)
    new_r = OxmlElement("w:r")
    if rPr is not None:
        new_r.append(copy.deepcopy(rPr))
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = new_text
    new_r.append(new_t)
    para._p.append(new_r)


def set_part_title(doc, part_no: str, part_name: str) -> None:
    """Header table, row 1, cell 2: '<part_no> - <part_name>'."""
    header = doc.sections[0].header
    if not header.tables:
        return
    para = header.tables[0].rows[1].cells[2].paragraphs[0]
    title = f"{part_no} - {part_name}" if part_no else part_name
    _replace_paragraph_text(para, title)


def set_doc_number(doc, doc_no: str) -> None:
    """Footer table, row 0, cell 0: '<doc_no>'."""
    footer = doc.sections[0].footer
    if not footer.tables:
        return
    para = footer.tables[0].rows[0].cells[0].paragraphs[0]
    _replace_paragraph_text(para, doc_no)


# ── Page table handling ─────────────────────────────────────────────────────

def _fill_page_table(table, page_steps, base_index: int) -> None:
    """Fill one 6-row table with up to 8 steps; blank unused headings.

    page_steps: list of up to 8 step dicts (may contain None / empty entries).
    base_index: step number of the first cell on this page minus 1
                (page 1 -> 0, page 2 -> 8, ...).
    """
    for slot in range(STEPS_PER_PAGE):
        block = slot // COLS          # 0 for steps 1-4, 1 for steps 5-8
        col = slot % COLS
        step_no = base_index + slot + 1

        step = page_steps[slot] if slot < len(page_steps) else None
        has_content = bool(step and (step.get("image") or step.get("text")))

        heading_cell = table.rows[HEADING_ROWS[block]].cells[col]
        image_cell = table.rows[IMAGE_ROWS[block]].cells[col]
        text_cell = table.rows[TEXT_ROWS[block]].cells[col]

        if has_content:
            set_heading_text(heading_cell, f"STEP {step_no}")
            if step.get("image"):
                add_image_to_cell(image_cell, step["image"])
            if step.get("text"):
                add_text_to_cell(text_cell, step["text"])
        else:
            # Empty step: clear heading, image, and text so copied content doesn't bleed through.
            set_heading_text(heading_cell, "")
            _clear_cell(image_cell)
            _clear_cell(text_cell)


def _append_page_break_and_table(doc, source_table):
    """Add a page break, then a deep copy of source_table, return the new table."""
    body = doc.element.body
    children = list(body)
    sectPr = body.find(qn("w:sectPr"))
    insert_idx = children.index(sectPr) if sectPr is not None else len(children)

    break_para = OxmlElement("w:p")
    run_el = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_el.append(br)
    break_para.append(run_el)
    body.insert(insert_idx, break_para)

    new_tbl = copy.deepcopy(source_table._tbl)
    body.insert(insert_idx + 1, new_tbl)
    from docx.table import Table
    return Table(new_tbl, doc)


# ── Public entry point ──────────────────────────────────────────────────────

def fill_sop(steps,
             part_no: str,
             part_name: str,
             doc_no: str,
             template_path: str,
             output_path: Optional[str] = None) -> Union[str, bytes]:
    """Fill the SOP template and save (or return) the .docx.

    Args:
        steps: list of {"image": path|bytes, "text": str}, any length.
        part_no, part_name, doc_no: header/footer values.
        template_path: path to SOP_Template_A3.docx (read-only source).
        output_path: where to write the .docx. If None, returns bytes.

    Returns:
        output_path if given, else the .docx as bytes.
    """
    doc = Document(template_path)
    first_table = doc.tables[0]

    # Page 1
    _fill_page_table(first_table, steps[:STEPS_PER_PAGE], base_index=0)

    # Additional pages for steps beyond 8
    page = 1
    while page * STEPS_PER_PAGE < len(steps):
        start = page * STEPS_PER_PAGE
        page_steps = steps[start:start + STEPS_PER_PAGE]
        new_table = _append_page_break_and_table(doc, first_table)
        _fill_page_table(new_table, page_steps, base_index=start)
        page += 1

    # Header / footer fields
    set_part_title(doc, part_no, part_name)
    set_doc_number(doc, doc_no)

    if output_path:
        doc.save(output_path)
        return output_path

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


if __name__ == "__main__":
    # Smoke test: fill with placeholder images if run directly.
    import tempfile
    from PIL import Image, ImageDraw

    tmp = tempfile.mkdtemp()
    demo_steps = []
    for i in range(6):
        p = os.path.join(tmp, f"s{i}.jpg")
        im = Image.new("RGB", (400, 300), (60 + i * 25, 90, 120))
        ImageDraw.Draw(im).text((180, 140), f"Step {i + 1}", fill="white")
        im.save(p)
        demo_steps.append({"image": p, "text": f"Description for step {i + 1}."})

    here = os.path.dirname(os.path.abspath(__file__))
    template = os.path.join(here, "SOP_Template_A3.docx")
    out = os.path.join(tmp, "demo_output.docx")
    fill_sop(demo_steps, "36611", "Tastic Luminate Heat Module", "A0866",
             template_path=template, output_path=out)
    print("Wrote", out)
