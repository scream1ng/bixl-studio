"""Deep inspect reference template XML for exact heading color + all template margins."""
import os
from docx import Document
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(HERE, "..", "docx", "sop", "source")

REF = os.path.join(TEMPLATE_DIR, "SOP_Template_A3_Landscape w Table.docx")
doc = Document(REF)
tbl = doc.tables[0]

print("=== HEADING ROW 0 — RAW RUN XML (col 0) ===")
cell = tbl.rows[0].cells[0]
for para in cell.paragraphs:
    print(f"  para align: {para.alignment}")
    for r in para._p.findall(qn("w:r")):
        rPr = r.find(qn("w:rPr"))
        if rPr is not None:
            color_el = rPr.find(qn("w:color"))
            sz_el    = rPr.find(qn("w:sz"))
            bold_el  = rPr.find(qn("w:b"))
            shd_el   = rPr.find(qn("w:shd"))
            font_el  = rPr.find(qn("w:rFonts"))
            print(f"    color: {color_el.get(qn('w:val')) if color_el is not None else 'not set'}")
            print(f"    sz:    {sz_el.get(qn('w:val')) if sz_el is not None else 'not set'} half-pts")
            print(f"    bold:  {'yes' if bold_el is not None else 'no'}")
            print(f"    shd:   {shd_el.get(qn('w:fill')) if shd_el is not None else 'none'}")
            if font_el is not None:
                print(f"    font:  ascii={font_el.get(qn('w:ascii'))} hAnsi={font_el.get(qn('w:hAnsi'))}")

print("\n=== HEADING ROW 0 — CELL tcPr FULL XML ===")
import lxml.etree as etree
cell = tbl.rows[0].cells[0]
tcPr = cell._tc.find(qn("w:tcPr"))
if tcPr is not None:
    print(etree.tostring(tcPr, pretty_print=True).decode())

print("\n=== HEADING ROW 0 — PARA pPr FULL XML ===")
para = tbl.rows[0].cells[0].paragraphs[0]
pPr = para._p.find(qn("w:pPr"))
if pPr is not None:
    print(etree.tostring(pPr, pretty_print=True).decode())

print("\n=== ALL TEMPLATE MARGINS ===")
TEMPLATES = {
    "a3-landscape": "SOP Template (A3 Landscape).docx",
    "a3-portrait":  "SOP Template (A3 Portrait).docx",
    "a4-landscape": "SOP Template (A4 Landscape).docx",
    "a4-portrait":  "SOP Template (A4 Portrait).docx",
}
for key, fname in TEMPLATES.items():
    d = Document(os.path.join(TEMPLATE_DIR, fname))
    s = d.sections[0]
    hd = s.header_distance.cm if s.header_distance else 0
    fd = s.footer_distance.cm if s.footer_distance else 0
    usable_h = s.page_height.cm - s.top_margin.cm - s.bottom_margin.cm
    print(f"\n  {key}:")
    print(f"    page:    {s.page_width.cm:.2f} x {s.page_height.cm:.2f} cm")
    print(f"    margins: T={s.top_margin.cm:.2f} B={s.bottom_margin.cm:.2f} L={s.left_margin.cm:.2f} R={s.right_margin.cm:.2f}")
    print(f"    header_distance={hd:.2f}  footer_distance={fd:.2f}")
    print(f"    usable body height: {usable_h:.2f} cm")
    print(f"    usable body width:  {s.page_width.cm - s.left_margin.cm - s.right_margin.cm:.2f} cm")
