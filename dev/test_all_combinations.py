"""
test_all_combinations.py
Generate all 16 combinations: 4 formats x steps-per-page 1/2/4/8.
Table always fills ~90% of page height. Grid, borders, heading style
matched exactly to reference template.
Run: python test_all_combinations.py
Output: test_output/
"""
import io
import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage, ImageDraw

HERE = os.path.dirname(os.path.abspath(__file__))
TEMPLATE_DIR = os.path.join(HERE, "..", "docx", "sop", "source")
OUT_DIR = os.path.join(HERE, "..", "docx", "sop", "templates")

TEMPLATES = {
    "a3-landscape": "SOP Template (A3 Landscape).docx",
    "a3-portrait":  "SOP Template (A3 Portrait).docx",
    "a4-landscape": "SOP Template (A4 Landscape).docx",
    "a4-portrait":  "SOP Template (A4 Portrait).docx",
}
STEPS_OPTIONS = [1, 2, 4, 8]

# Reference block height from A3 landscape template (cm per block)
REF_HEADING_CM = 1.04
REF_IMAGE_CM   = 5.84
REF_TEXT_CM    = 3.05
REF_BLOCK_CM   = REF_HEADING_CM + REF_IMAGE_CM + REF_TEXT_CM  # 9.93 cm

# Target fill: table occupies this fraction of usable body height
TARGET_FILL = 0.90

# Image cell vertical padding (twips) — gap above and below picture
IMG_PAD_TWIPS = 100   # ~1.76 mm each side

# Grid: (steps_per_page, is_landscape) -> (cols, blocks)
# 4 steps = 2x2 on both orientations
GRID = {
    (1, True):  (1, 1),
    (2, True):  (2, 1),
    (4, True):  (2, 2),
    (8, True):  (4, 2),
    (1, False): (1, 1),
    (2, False): (1, 2),
    (4, False): (2, 2),
    (8, False): (2, 4),
}


def get_grid(is_landscape: bool, steps_per_page: int):
    return GRID[(steps_per_page, is_landscape)]


def get_row_heights_twips(usable_h_cm: float, n_blocks: int):
    """Scale reference heights so n_blocks fills TARGET_FILL of usable_h_cm."""
    target_cm = usable_h_cm * TARGET_FILL
    scale = target_cm / (n_blocks * REF_BLOCK_CM)
    def cm_to_twips(cm): return int(cm / 2.54 * 1440)
    h_h = cm_to_twips(REF_HEADING_CM * scale)
    i_h = cm_to_twips(REF_IMAGE_CM   * scale)
    t_h = cm_to_twips(REF_TEXT_CM    * scale)
    return h_h, i_h, t_h


# ── XML helpers ─────────────────────────────────────────────────────────────

def _set_row_height(row, twips: int) -> None:
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), str(twips))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)


def _set_cell_bg(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_valign(cell, align: str = "center") -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:vAlign")):
        tcPr.remove(old)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), align)
    tcPr.append(va)


def _set_cell_padding(cell, top=0, bottom=0, left=115, right=115) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_col_width(cell, twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcW")):
        tcPr.remove(old)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _set_table_props(table, width_twips: int, ind_twips: int = 0) -> None:
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Width
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(width_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

    # Left-aligned with explicit indent (centers on page even with asymmetric margins)
    for old in tblPr.findall(qn("w:jc")):
        tblPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    tblPr.append(jc)

    for old in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(old)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(ind_twips))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    # Borders: 3pt solid black (matched to reference)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "24")        # 3 pt — matches reference
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _clear_cell(cell):
    for p in cell.paragraphs[1:]:
        p._element.getparent().remove(p._element)
    for r in list(cell.paragraphs[0].runs):
        r._element.getparent().remove(r._element)


# ── Content writers ──────────────────────────────────────────────────────────

def write_heading(cell, text: str) -> None:
    """Reference style: white bg, Century Gothic 14pt bold underlined, 1.5x spacing."""
    _set_cell_bg(cell, "FFFFFF")
    _set_cell_valign(cell, "center")
    _set_cell_padding(cell, top=0, bottom=0, left=115, right=115)
    _clear_cell(cell)
    para = cell.paragraphs[0]
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "360")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = "Century Gothic"
    run.font.size = Pt(14)
    run.font.bold = True
    rPr = run._r.get_or_add_rPr()
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)


def write_image(cell, img_bytes: bytes, col_w_twips: int, img_h_twips: int) -> None:
    pad_cm = IMG_PAD_TWIPS / 1440 * 2.54
    col_w_cm = col_w_twips / 1440 * 2.54
    img_h_cm = (img_h_twips / 1440 * 2.54) - (pad_cm * 2)
    _set_cell_bg(cell, "FFFFFF")
    _set_cell_valign(cell, "center")
    _set_cell_padding(cell, top=IMG_PAD_TWIPS, bottom=IMG_PAD_TWIPS, left=115, right=115)
    _clear_cell(cell)
    pil = PILImage.open(io.BytesIO(img_bytes))
    w, h = pil.size
    px_per_cm = 37.795
    scale = min((col_w_cm * px_per_cm) / w, (img_h_cm * px_per_cm) / h)
    final_w_cm = (w * scale) / px_per_cm
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(io.BytesIO(img_bytes), width=Cm(final_w_cm))


def write_text(cell, text: str) -> None:
    _set_cell_bg(cell, "FFFFFF")
    _set_cell_valign(cell, "top")
    _set_cell_padding(cell, top=80, bottom=0, left=115, right=115)
    _clear_cell(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    run.font.name = "Century Gothic"
    run.font.size = Pt(11)


# ── Table builder ────────────────────────────────────────────────────────────

def build_table(doc, steps: list, cols: int, n_blocks: int,
                usable_w_cm: float, usable_h_cm: float, ind_cm: float = 0.0) -> None:
    total_rows = n_blocks * 3
    h_h, i_h, t_h = get_row_heights_twips(usable_h_cm, n_blocks)
    usable_w_twips = int(usable_w_cm / 2.54 * 1440)
    ind_twips = int(ind_cm / 2.54 * 1440)
    col_w_twips = usable_w_twips // cols

    table = doc.add_table(rows=total_rows, cols=cols)
    _set_table_props(table, usable_w_twips, ind_twips)

    for col in range(cols):
        _set_col_width(table.rows[0].cells[col], col_w_twips)

    for block in range(n_blocks):
        hr = block * 3
        ir = block * 3 + 1
        tr = block * 3 + 2
        _set_row_height(table.rows[hr], h_h)
        _set_row_height(table.rows[ir], i_h)
        _set_row_height(table.rows[tr], t_h)

        for col in range(cols):
            slot = block * cols + col
            step = steps[slot] if slot < len(steps) else None
            has = bool(step and (step.get("image") or step.get("text")))

            hcell = table.rows[hr].cells[col]
            icell = table.rows[ir].cells[col]
            tcell = table.rows[tr].cells[col]

            if has:
                write_heading(hcell, f"STEP {slot + 1}")
                if step.get("image"):
                    write_image(icell, step["image"], col_w_twips, i_h)
                else:
                    _set_cell_bg(icell, "F2F2F2")
                write_text(tcell, step.get("text", ""))
            else:
                write_heading(hcell, "")
                _set_cell_bg(icell, "F2F2F2")
                _set_cell_bg(tcell, "FFFFFF")


# ── Placeholder steps ────────────────────────────────────────────────────────

COLORS = [
    (180,60,60),(60,120,180),(60,160,80),(160,100,40),
    (100,60,160),(40,160,160),(180,130,40),(80,80,80),
]
DESCS = [
    "Remove the four M4 retaining screws from the front housing.",
    "Lift the heat module clear, keeping the ribbon cable attached.",
    "Disconnect the ribbon cable at the board connector.",
    "Inspect the thermal pad for damage or displacement.",
    "Position new module; align the four mounting holes.",
    "Reconnect the ribbon cable — ensure the latch clicks.",
    "Lower the module onto the housing and hand-start screws.",
    "Torque screws to 1.2 Nm in a cross pattern.",
]


def _make_img(color, step_no: int) -> bytes:
    img = PILImage.new("RGB", (600, 400), color)
    draw = ImageDraw.Draw(img)
    draw.rectangle([20, 20, 579, 379], outline="white", width=3)
    draw.text((260, 185), f"Step {step_no}", fill="white")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    return buf.getvalue()


def make_steps(n: int) -> list:
    return [
        {"image": _make_img(COLORS[i % 8], i + 1), "text": DESCS[i % 8]}
        for i in range(n)
    ]


# ── Generator ────────────────────────────────────────────────────────────────

def generate(template_key: str, steps_per_page: int) -> str:
    tpl_path = os.path.join(TEMPLATE_DIR, TEMPLATES[template_key])
    doc = Document(tpl_path)
    section = doc.sections[0]

    page_w_cm = section.page_width.cm
    page_h_cm = section.page_height.cm
    ml = section.left_margin.cm   if section.left_margin   else 2.54
    mr = section.right_margin.cm  if section.right_margin  else 1.50
    mt = section.top_margin.cm    if section.top_margin    else 2.54
    mb = section.bottom_margin.cm if section.bottom_margin else 2.54

    # Use symmetric margin (larger side) so table is visually centered on page.
    # tbl_ind shifts table right by the difference when left margin > right margin.
    sym = max(ml, mr)
    usable_w = page_w_cm - 2 * sym
    tbl_ind = sym - ml          # always >= 0; 0 when margins already symmetric
    usable_h = page_h_cm - mt - mb

    is_landscape = page_w_cm > page_h_cm
    cols, n_blocks = get_grid(is_landscape, steps_per_page)

    # Landscape: PPE is in Word header — body is only empty paragraphs, safe to strip.
    # Portrait: PPE table + list are in body — preserve them, append step table at end.
    if is_landscape:
        body = doc.element.body
        for child in list(body):
            tag = child.tag.split("}")[-1]
            if tag in ("tbl", "p"):
                body.remove(child)

    steps = make_steps(steps_per_page)
    build_table(doc, steps, cols, n_blocks, usable_w, usable_h, tbl_ind)

    fname = f"{template_key}__{steps_per_page}steps.docx"
    out = os.path.join(OUT_DIR, fname)
    doc.save(out)
    return fname


def main():
    os.makedirs(OUT_DIR, exist_ok=True)
    for tpl_key in TEMPLATES:
        for spp in STEPS_OPTIONS:
            fname = generate(tpl_key, spp)
            print(f"  {fname}")
    print(f"\n16 files -> {OUT_DIR}")


if __name__ == "__main__":
    main()
